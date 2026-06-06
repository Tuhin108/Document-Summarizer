"""
document_parser.py — Multi-format Document Text Extractor
Supports PDF (via PyMuPDF), DOCX (via python-docx), TXT, MD, CSV.
"""

import io
import csv
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentParser:
    """Routes file bytes to the appropriate parser based on MIME type."""

    MAX_CHARS = 15_000  # Cap text sent to the AI to manage tokens

    def extract_text(self, content: bytes, mime_type: str, filename: str) -> Optional[str]:
        """
        Main entry point. Returns plain text or None on failure.
        mime_type is used first; filename extension as fallback.
        """
        mime_type = mime_type.lower()
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        try:
            if mime_type == "application/pdf" or ext == "pdf":
                text = self._parse_pdf(content)
            elif (
                mime_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                or ext in ("docx",)
            ):
                text = self._parse_docx(content)
            elif mime_type in ("application/msword",) or ext == "doc":
                # Attempt DOCX path (works if file is actually DOCX-compatible)
                text = self._parse_docx(content)
            elif mime_type in ("text/plain", "text/markdown") or ext in ("txt", "md"):
                text = self._parse_text(content)
            elif mime_type == "text/csv" or ext == "csv":
                text = self._parse_csv(content)
            else:
                # Last-resort: try plain-text decode
                text = self._parse_text(content)

            if text:
                text = self._clean(text)
                return text[: self.MAX_CHARS]
            return None

        except Exception as exc:
            logger.error("Parser error for '%s': %s", filename, exc)
            raise

    # ────────────────────────────────────────────
    # Format-specific parsers
    # ────────────────────────────────────────────

    def _parse_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError(
                "PyMuPDF not installed. Run: pip install PyMuPDF"
            )

        pages = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page in doc:
                pages.append(page.get_text("text"))

        return "\n\n".join(pages)

    def _parse_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx not installed. Run: pip install python-docx"
            )

        doc = Document(io.BytesIO(content))
        parts = []

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    parts.append(" | ".join(row_cells))

        return "\n\n".join(parts)

    def _parse_text(self, content: bytes) -> str:
        """Decode raw text with encoding auto-detection."""
        for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace")

    def _parse_csv(self, content: bytes) -> str:
        """Convert CSV rows into readable prose."""
        text = self._parse_text(content)
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return ""

        headers = rows[0]
        lines = [" | ".join(headers), "-" * 40]
        for row in rows[1:31]:  # first 30 data rows
            lines.append(" | ".join(row))

        if len(rows) > 31:
            lines.append(f"... and {len(rows) - 31} more rows")

        return "\n".join(lines)

    # ────────────────────────────────────────────
    # Text cleaning
    # ────────────────────────────────────────────

    def _clean(self, text: str) -> str:
        """Normalise whitespace and remove junk characters."""
        import re

        # Collapse excessive newlines
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Collapse excessive spaces
        text = re.sub(r"[ \t]{2,}", " ", text)
        # Strip non-printable characters (keep newlines)
        text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", "", text)
        return text.strip()
